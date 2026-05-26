import docx
import re

def clean_text(text):
    replacements = [
        (r'\bкритичен\b', 'важен'),
        (r'\bКритичен\b', 'Важен'),
        (r'\bкритична\b', 'важна'),
        (r'\bКритична\b', 'Важна'),
        (r'\bкритично\b', 'важно'),
        (r'\bКритично\b', 'Важно'),
        (r'\bВ заключение,?\s*', ''),
        (r'\bВажно е да се отбележи, че\s*', ''),
        (r'\bСтрува си да се спомене, че\s*', ''),
        (r'\bВ обобщение,?\s*', ''),
        (r'\bОт съществено значение е\b', 'Необходимо е'),
        (r'\bключов\b', 'основен'),
        (r'\bКлючов\b', 'Основен'),
        (r'\bключова\b', 'основна'),
        (r'\bКлючова\b', 'Основна'),
        (r'\bключово\b', 'основно'),
        (r'\bКлючово\b', 'Основно'),
        (r'\bжизненоважно\b', 'важно'),
        (r'\bЖизненоважно\b', 'Важно'),
        (r'\bжизненоважен\b', 'важен'),
        (r'\bЖизненоважен\b', 'Важен'),
        (r'\bжизненоважна\b', 'важна'),
        (r'\bЖизненоважна\b', 'Важна'),
        (r'\bизключително\b', 'много'),
        (r'\bИзключително\b', 'Много'),
        (r'\bТози комплексен подход\b', 'Този подход')
    ]
    for old, new in replacements:
        text = re.sub(old, new, text)
    return text

def is_code(text):
    text = text.strip()
    if not text:
        return False
    if text.startswith(('# ', '//', '/*', '*/', 'import ', 'export ', 'def ', 'class ', 'const ', 'let ', '<', 'ST_', '@', 'module.exports')):
        return True
    if text.endswith(('{', '}', ';', '/>', '",', "',")):
        return True
    if ' = ' in text and ('(' in text or '[' in text or '{' in text):
        return True
    return False

def balance_docx(path):
    doc = docx.Document(path)
    
    # 1. Clean parasitic phrases
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        original = p.text
        cleaned = clean_text(original)
        if cleaned != original:
            # We replace text. To preserve style somewhat, we just assign to p.text
            # Note: This removes inline runs formatting (like partial bold), but it's acceptable for simple words.
            # To be safer, we can just replace text in runs if the phrase doesn't cross run boundaries.
            # But regex sub on runs is tricky. Since this is an AI edit, replacing p.text is standard.
            style = p.style
            p.text = cleaned
            p.style = style

    # 2. Balance code blocks
    # Group paragraphs into chapters (1. 2. 3. 4.)
    chapters = []
    current_chapter = None
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if text and text[0].isdigit() and "." in text[:3] and len(text) < 100 and text[1:3] in (". ", ".  ", ".\t", " ."):
            current_chapter = {"name": text, "paragraphs": [], "code_blocks": []}
            chapters.append(current_chapter)
        elif current_chapter:
            current_chapter["paragraphs"].append(p)
            
    # If no strict chapters found (e.g. they don't match the heuristic), fallback to one block
    if not chapters:
        chapters = [{"name": "All", "paragraphs": doc.paragraphs, "code_blocks": []}]
        
    for ch in chapters:
        # Group paragraphs into code blocks
        blocks = []
        current_block = []
        for p in ch["paragraphs"]:
            if is_code(p.text):
                current_block.append(p)
            else:
                if current_block:
                    blocks.append(current_block)
                    current_block = []
        if current_block:
            blocks.append(current_block)
            
        ch["code_blocks"] = blocks
        
    # We want max ~8-10 code blocks per chapter, and max ~6 lines per code block
    target_blocks_per_chapter = 8
    target_lines_per_block = 5
    
    for ch in chapters:
        blocks = ch["code_blocks"]
        
        # Trim blocks that are too long
        for b in blocks:
            if len(b) > target_lines_per_block:
                # keep first few lines, add a comment line "... (truncated)", remove the rest
                for p_idx in range(target_lines_per_block - 1, len(b)):
                    if p_idx == target_lines_per_block - 1:
                        b[p_idx].text = "    // ... "
                    else:
                        p_to_remove = b[p_idx]
                        p_to_remove._element.getparent().remove(p_to_remove._element)
                        
        # Drop excessive blocks
        if len(blocks) > target_blocks_per_chapter:
            # We need to drop some blocks. Let's keep the first few and last few, drop middle ones.
            num_to_drop = len(blocks) - target_blocks_per_chapter
            # Drop every Nth block until we reach target
            blocks_to_drop = []
            step = max(1, len(blocks) // num_to_drop)
            for i in range(1, len(blocks)-1, step):
                if len(blocks_to_drop) < num_to_drop:
                    blocks_to_drop.append(blocks[i])
            
            for b in blocks_to_drop:
                for p in b:
                    # check if already removed
                    if p._element.getparent() is not None:
                        p._element.getparent().remove(p._element)

    doc.save(path)
    print("Document successfully balanced and cleaned.")

if __name__ == "__main__":
    balance_docx(r"C:\Users\irena\Downloads\GLAVA_4.docx")
