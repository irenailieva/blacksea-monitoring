import docx

def process_doc():
    doc = docx.Document(r"C:\Users\irena\Downloads\GLAVA_4_backup.docx")
    
    blocks = [
        {
            "match": "def preprocess_raster",
            "replacement": '''# Функция за предварителна обработка (preprocessing) на растерно изображение
def preprocess_raster(input_path: str) -> str:
    """
    Preprocesses the input raster.
    Currently implements basic normalization or pass-through.
    """
    logger.info(f"Preprocessing {input_path}...")
    
    # В момента функцията създава копие на изображението.
    # В реални сценарии тук се прилагат алгоритми за атмосферна корекция (напр. Sen2Cor)
    # или маскиране на облаци.
    
    # Дефиниране на пътя за изходния (обработен) файл
    output_path = input_path.replace(".tif", "_processed.tif")
    
    # Отваряне на суровия растерен файл
    with rasterio.open(input_path) as src:
        # Извличане на метаданните (георефериране, брой канали, тип данни)
        profile = src.profile
        # Изчитане на всички пиксели в паметта
        data = src.read()
        
        # Отваряне на новия файл и записване на обработените данни
        with rasterio.open(output_path, 'w', **profile) as dst:
            dst.write(data)
            
            # Копиране на описанията (имената) на всеки канал от оригиналния към новия файл
            for i in range(1, src.count + 1):
                dst.set_band_description(i, src.descriptions[i-1])
                
    logger.success(f"Preprocessing complete. Saved to {output_path}")
    return output_path'''
        },
        {
            "match": "def generate_index",
            "replacement": '''# Функция за генериране на спектрални индекси (например NDVI) от предварително обработени растерни данни
def generate_index(raster_path: str, index_name: str = 'NDVI') -> str:
    """
    Generates a spectral index from the input raster.
    """
    logger.info(f"Generating {index_name} for {raster_path}...")
    
    # Формиране на пътя за изходния файл, добавяйки името на индекса като суфикс
    output_path = raster_path.replace(".tif", f"_{index_name}.tif")
    
    # Отваряне на растерния файл за четене
    with rasterio.open(raster_path) as src:
        # Извличане на Червения (Red) и Близкия инфрачервен (NIR) канал.
        # В Sentinel-2 данните: B4 е Червеният канал (индекс 3), B8 е NIR (индекс 4).
        red = src.read(3).astype('float32')
        nir = src.read(4).astype('float32')
        
        # Защита срещу деление на нула чрез добавяне на много малко число (epsilon)
        epsilon = 1e-10
        # Пресмятане на Нормализиран диференциален вегетационен индекс (NDVI):
        ndvi = (nir - red) / (nir + red + epsilon)
        
        # Копиране на метаданните (profile) от оригиналния файл
        profile = src.profile
        profile.update(
            dtype=rasterio.float32,
            count=1,
            driver='GTiff'
        )
        
        # Отваряне на нов файл за запис (write) с обновените метаданни
        with rasterio.open(output_path, 'w', **profile) as dst:
            # Запис на изчисления индекс като първи канал (band 1)
            dst.write(ndvi, 1)
            # Задаване на описание (име) на канала
            dst.set_band_description(1, index_name)
            
    logger.success(f"{index_name} generated at {output_path}")
    return output_path'''
        },
        {
            "match": "def update_job_status",
            "replacement": '''# Функция за обновяване на статуса на дадена задача (Job) в базата данни
def update_job_status(engine, job_id, status, progress=None):
    if not job_id:
        return
    try:
        # Отваряне на транзакция към базата данни
        with engine.begin() as conn:
            # Ако е подаден прогрес, обновяваме JSON полето payload
            if progress is not None:
                progress_json = json.dumps(progress)
                prog_str = (
                    f", payload = jsonb_set(COALESCE(payload::jsonb, '{{}}'::jsonb), "
                    f"'{{progress}}', '{progress_json}'::jsonb)::json "
                )
            else:
                prog_str = ""
            
            # Ако задачата е приключила (успешно или не), записваме часа на приключване
            finish_str = ", finished_at = NOW() " if status in ('completed', 'failed') else ""
            
            # Подготовка и изпълнение на SQL заявката за обновяване
            stmt = (
                f"UPDATE etl_jobs SET status = :status, updated_at = NOW() "
                f"{prog_str}{finish_str}WHERE id = :job_id"
            )
            conn.execute(text(stmt), {"status": status, "job_id": job_id})
            logger.info(f"Job {job_id} → {status}" + (f" ({progress}%)" if progress is not None else ""))
    except Exception as e:
        logger.error(f"Failed to update job status for job {job_id}: {e}")'''
        },
        {
            "match": "rf = RandomForestClassifier",
            "replacement": '''    # 1. Random Forest Classifier
    # Този модел създава множество "дървета на решенията" и осреднява техните резултати.
    # Много добър за предотвратяване на преобучаване (overfitting).
    print("Обучение на Random Forest...")
    rf = RandomForestClassifier(n_estimators=50, max_depth=10, n_jobs=-1, class_weight='balanced', random_state=RANDOM_SEED)
    rf.fit(X_train, y_train)'''
        },
        {
            "match": "xgb_model = xgb.XGBClassifier",
            "replacement": '''    # 2. XGBoost (eXtreme Gradient Boosting)
    # Много силен алгоритъм, базиран на последователно надграждане на слаби дървета.
    print("Обучение на XGBoost...")
    xgb_model = xgb.XGBClassifier(
        n_estimators=50, max_depth=6, objective='multi:softmax', num_class=3,
        tree_method="hist", n_jobs=-1, random_state=RANDOM_SEED
    )
    xgb_model.fit(X_train, y_train)'''
        },
        {
            "match": "lgb_model = lgb.LGBMClassifier",
            "replacement": '''    # 3. LightGBM (Light Gradient Boosting Machine)
    # Подобен на XGBoost, но обикновено по-бърз и по-ефективен при много данни.
    print("Обучение на LightGBM...")
    lgb_model = lgb.LGBMClassifier(
        n_estimators=50, num_leaves=31, objective='multiclass', 
        class_weight='balanced', verbose=-1, random_state=RANDOM_SEED, n_jobs=-1
    )
    lgb_model.fit(X_train, y_train)'''
        },
        {
            "match": "X_train, X_test, y_train, y_test = train_test_split",
            "replacement": '''    # Разделяне на данните на тренировъчни (80%) и тестови (20%)
    # Параметърът stratify=y гарантира, че съотношението на класовете ще е еднакво и в двете множества.
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=RANDOM_SEED, stratify=y)'''
        },
        {
            "match": "def _vote(self, X):",
            "replacement": '''    def _vote(self, X):
        """
        Вътрешна функция за осъществяване на Soft Voting (меко гласуване).
        Всеки от трите модела дава вероятност за всеки клас. Тези вероятности
        се усредняват, и печели класът с най-висока средна вероятност.
        Формата на входа X трябва да е (брои_примери, брой_признаци).
        """
        # Извличане на вероятностите (predict_proba)
        prob1 = self.rf.predict_proba(X)
        prob2 = self.xgb.predict_proba(X)
        prob3 = self.lgb.predict_proba(X)
        
        # Усредняване на вероятностите
        avg_prob = (prob1 + prob2 + prob3) / 3.0
        
        # Избиране на класа с най-висока усреднена вероятност
        final_pred = np.argmax(avg_prob, axis=1)
        return final_pred'''
        },
        {
            "match": "@app.post(\"/infer_batch\"",
            "replacement": '''@app.post("/infer_batch", response_model=PredictBatchResponse)
def infer_batch(req: PredictBatchRequest):
    """
    Endpoint за масово предсказване на много точки едновременно.
    Използва се за по-висока производителност.
    """
    if not model:
        raise HTTPException(status_code=503, detail="Моделът не е зареден")

    try:
        yhat = model.predict_batch(req.batch)
        return PredictBatchResponse(yhat=yhat)
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))'''
        }
    ]

    garbage_prefixes = [
        "return output_path", "# Връщане", "# ...", "return result",
        "payload = jsonb_set", "updated_at = NOW()",
        "class_weight=", "tree_method=", "objective=",
        "prob1 =", "prob2 =", "prob3 =", "# Усредняване", "avg_prob =",
        "# Избор на клас", "final_pred =", "return final_pred",
        "yhat =", "return PredictBatchResponse", "raise HTTPException",
        "}", ");", "])", "...", "]", "# Дефиниране", "# Изпълнение"
    ]
    
    i = 0
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        text = p.text.strip()
        
        for block in blocks:
            if block["match"] in text:
                print(f"Matched {block['match']}")
                # Replace the text. We can set it in the first run to keep paragraph styling roughly,
                # or just set p.text. Let's set p.text.
                p.text = block["replacement"]
                
                # Check subsequent paragraphs to clear garbage
                j = i + 1
                while j < len(doc.paragraphs):
                    next_text = doc.paragraphs[j].text.strip()
                    if next_text == "":
                        # Ignore empty lines
                        doc.paragraphs[j].text = ""
                        j += 1
                        continue
                        
                    is_garbage = False
                    for pfx in garbage_prefixes:
                        if next_text.startswith(pfx) or pfx in next_text:
                            is_garbage = True
                            break
                    
                    if is_garbage:
                        doc.paragraphs[j].text = ""
                        j += 1
                    else:
                        break
                i = j - 1
                break
        i += 1
        
    doc.save(r"C:\Users\irena\Downloads\GLAVA_4.docx")
    print("Done writing to GLAVA_4.docx")

if __name__ == "__main__":
    process_doc()
